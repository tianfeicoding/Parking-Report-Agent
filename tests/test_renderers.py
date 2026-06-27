"""图表和 DOCX 渲染测试。

本文件验证 Phase 7 的输出是真实 PNG 图表和可打开的 Word 文档，
并检查报告中包含关键硬指标。
"""

from decimal import Decimal
from zipfile import ZipFile

from docx import Document

from app.agent.report_plan_schema import ChartPlan, Observation, Recommendation, ReportPlan
from app.data.parking_csv_loader import load_parking_csv
from app.metrics.parking_metrics import compute_parking_metrics
from app.profiling.duration_profiler import build_duration_profile
from app.profiling.payment_profiler import build_payment_profile
from app.report.chart_renderer import ChartRenderer
from app.report.docx_renderer import render_report_docx
from app.report.template_instructions import TemplateInstructions


CSV_HEADER = (
    "应收金额,实收金额(元),免费金额(元),充值卡扣费(元),抵扣金额(元),"
    "抵扣时长(小时),实际抵扣额(元),支付方式,支付渠道,收费时间,进车时间\n"
)


def write_csv(tmp_path):
    path = tmp_path / "data.csv"
    path.write_text(
        CSV_HEADER
        + "30,30,0,0.00,0,0,0,微信,线上支付,2026-04-30 10:00:00,2026-04-30 08:00:00\n"
        + "45,0,0,0.00,45,0,45,会员积分,线上支付,2026-04-30 13:30:00,2026-04-30 09:00:00\n"
        + "25,20,0,0.00,5,1,5,微信,出口贴码,2026-04-30 23:00:00,2026-04-30 10:00:00\n",
        encoding="utf-8-sig",
    )
    return path


def build_plan() -> ReportPlan:
    return ReportPlan(
        selected_charts=[
            ChartPlan(
                chart_id="payment_method_distribution",
                title="支付方式分布",
                reason="支撑支付结构分析",
                source_fact_ids=["payment_profile.dominant_payment_method"],
            )
        ],
        payment_section_summary="支付方式以微信为主，需要结合会员积分观察抵扣影响。",
        duration_section_summary="停车时长以短停为主，但存在超过 12 小时记录。",
        observations=[
            Observation(
                id="obs_payment",
                title="支付结构集中",
                text="支付方式集中在少数方式。",
                business_implication="可优先复核高频方式。",
                source_fact_ids=["payment_profile.dominant_payment_method"],
            ),
            Observation(
                id="obs_duration",
                title="存在长停记录",
                text="存在超过 12 小时的停车记录。",
                business_implication="需区分真实长停和记录异常。",
                source_fact_ids=["duration.over_12h_count"],
            ),
        ],
        recommendations=[
            Recommendation(
                id="rec_payment",
                text="按支付方式定期复核优惠和抵扣规则。",
                source_fact_ids=["payment_profile.dominant_payment_method"],
                source_observation_ids=["obs_payment"],
            ),
            Recommendation(
                id="rec_duration",
                text="抽样复核长时停车记录。",
                source_fact_ids=["duration.over_12h_count"],
                source_observation_ids=["obs_duration"],
            ),
        ],
    )


def test_chart_renderer_generates_all_phase7_charts(tmp_path):
    data = load_parking_csv(write_csv(tmp_path))
    payment_profile = build_payment_profile(data)
    duration_profile = build_duration_profile(data)

    charts = ChartRenderer(tmp_path / "charts").render_all(payment_profile, duration_profile)

    assert set(charts) == {
        "payment_method_distribution",
        "payment_channel_distribution",
        "parking_duration_distribution",
        "entry_hour_distribution",
        "charge_hour_distribution",
    }
    assert {chart.title for chart in charts.values()} == {
        "支付方式分布",
        "支付渠道分布",
        "停车时长分布",
        "入场时段分布",
        "收费时段分布",
    }
    for chart in charts.values():
        assert chart.path.exists()
        assert chart.path.stat().st_size > 0


def test_docx_renderer_writes_metrics_and_real_chart(tmp_path):
    data = load_parking_csv(write_csv(tmp_path))
    metrics = compute_parking_metrics(data)
    payment_profile = build_payment_profile(data)
    duration_profile = build_duration_profile(data)
    charts = ChartRenderer(tmp_path / "charts").render_all(payment_profile, duration_profile)
    output_path = tmp_path / "report.docx"

    render_report_docx(
        output_path=output_path,
        template_instructions=TemplateInstructions(
            title="停车明细分析报告",
            sections=["一、关键指标"],
            placeholders=["【填写指标】"],
            raw_text="停车明细分析报告\n一、关键指标\n【填写指标】",
        ),
        metrics=metrics,
        payment_profile=payment_profile,
        duration_profile=duration_profile,
        report_plan=build_plan(),
        charts=charts,
    )

    assert output_path.exists()
    with ZipFile(output_path) as docx:
        names = set(docx.namelist())
        document_xml = docx.read("word/document.xml").decode("utf-8")
        styles_xml = docx.read("word/styles.xml").decode("utf-8")

    assert "word/media/image1.png" in names
    assert "交易总笔数" in document_xml
    assert "100.00" in document_xml
    assert "实收率" in document_xml
    assert f"{Decimal('50.0')}%" in document_xml
    assert "【填写指标】" not in document_xml
    assert "Microsoft YaHei" in styles_xml


def write_template_docx(tmp_path):
    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_heading("停车明细分析报告", level=0)
    doc.add_paragraph("【报告模板 —— 红色斜体为需由智能体填写的占位内容】")
    doc.add_heading("报告信息", level=2)
    doc.add_paragraph("数据周期：【起始日期 – 结束日期】")
    doc.add_paragraph("生成时间：【时间戳】")
    doc.add_heading("一、关键指标（硬性数字）", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "填写值"
    table.rows[0].cells[2].text = "示例预期值"
    rows = [
        ("1. 总交易笔数", "[ 笔数 ]", "= 14"),
        ("2. 应收总金额（元）", "[ ¥ ___ ]", "= ¥540"),
        ("3. 实收总金额（元）", "[ ¥ ___ ]", "= ¥390"),
        ("4. 实际抵扣总额（元）", "[ ¥ ___ ]", "= ¥150"),
        ("5. 实收率（%）", "[ __._% ]", "= 72.2%"),
        ("6. 主要支付方式", "[ ____ ×N ]", "= 微信 ×6"),
    ]
    for label, placeholder, expected in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = placeholder
        cells[2].text = expected
    doc.add_heading("二、支付方式与渠道", level=1)
    doc.add_paragraph("【叙述：各支付方式与支付渠道（线上支付 / 出口贴码）的分布情况，指出占主导的方式与渠道。】")
    doc.add_paragraph("【占位图 —— 须由智能体替换为依据实际数据生成的图表。报告至少须包含一张真实图表。】")
    doc.add_heading("三、停车时长分析", level=1)
    doc.add_paragraph("【根据“收费时间 − 进车时间”推算停车时长。报告平均值 / 分布情况及明显的长时停车异常项。")
    doc.add_paragraph("停车时长分布柱形图（停车小时数，每两小时递增，0-12小时）")
    doc.add_paragraph("】")
    doc.add_heading("四、补充观察", level=1)
    doc.add_paragraph("【智能体生成的观察内容置于此处 —— 2 至 3 条简洁、有数据支撑、与管理决策相关的要点。")
    doc.add_paragraph("读者对象：负责经营本停车业务的管理者。")
    doc.add_heading("五、结论与建议", level=1)
    doc.add_paragraph("【依据上述分析给出 2 至 4 条可执行的建议。如无建议，请说明。】")
    doc.save(path)
    return path


def test_docx_renderer_removes_example_column_and_replaces_placeholders(tmp_path):
    data = load_parking_csv(write_csv(tmp_path))
    metrics = compute_parking_metrics(data)
    payment_profile = build_payment_profile(data)
    duration_profile = build_duration_profile(data)
    charts = ChartRenderer(tmp_path / "charts").render_all(payment_profile, duration_profile)
    template_path = write_template_docx(tmp_path)
    output_path = tmp_path / "templated-report.docx"

    render_report_docx(
        output_path=output_path,
        template_path=template_path,
        template_instructions=TemplateInstructions(
            title="停车明细分析报告",
            sections=[],
            placeholders=[],
            raw_text="",
        ),
        metrics=metrics,
        payment_profile=payment_profile,
        duration_profile=duration_profile,
        report_plan=build_plan(),
        charts=charts,
        data_period="2026-04-30",
    )

    rendered = Document(output_path)
    metrics_table = rendered.tables[0]
    assert len(metrics_table.columns) == 2
    assert metrics_table.rows[0].cells[0].text == "指标"
    assert metrics_table.rows[0].cells[1].text == "填写值"
    assert metrics_table.rows[1].cells[1].text == "3"
    assert metrics_table.rows[2].cells[1].text == "¥100.00"
    assert metrics_table.rows[6].cells[1].text == "微信 ×2"

    with ZipFile(output_path) as docx:
        names = set(docx.namelist())
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "word/media/image1.png" in names
    assert "示例预期值" not in document_xml
    assert "= 14" not in document_xml
    assert "【占位图" not in document_xml
    assert "[ 笔数 ]" not in document_xml
    assert "【智能体生成" not in document_xml
