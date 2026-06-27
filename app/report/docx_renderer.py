"""Word 报告渲染器。

本文件把硬指标、Agent ReportPlan、profile 摘要和真实图表写入 .docx。
渲染器优先在用户上传的原模板上替换占位内容，保留模板版式；
渲染器不让 LLM 改写硬指标，只消费已经校验过的结构化输入。
"""

from collections.abc import Callable
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from app.agent.report_plan_schema import ReportPlan
from app.metrics.parking_metrics import ParkingMetrics
from app.profiling.duration_profiler import DurationProfile
from app.profiling.payment_profiler import PaymentProfile
from app.report.chart_renderer import RenderedChart
from app.report.template_instructions import TemplateInstructions


BODY_FONT = "Microsoft YaHei"
HEADING_FONT = "Microsoft YaHei"


def render_report_docx(
    *,
    output_path: Path,
    template_instructions: TemplateInstructions,
    metrics: ParkingMetrics,
    payment_profile: PaymentProfile,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
    template_path: str | Path | None = None,
    data_period: str | None = None,
    generated_at: datetime | None = None,
    on_template_fallback: Callable[[str], None] | None = None,
) -> Path:
    """生成最终 Word 报告，并返回输出路径。"""
    if template_path is not None:
        try:
            return _render_template_aware_report(
                output_path=output_path,
                template_path=Path(template_path),
                metrics=metrics,
                payment_profile=payment_profile,
                duration_profile=duration_profile,
                report_plan=report_plan,
                charts=charts,
                data_period=data_period,
                generated_at=generated_at,
            )
        except Exception as exc:
            if on_template_fallback is not None:
                on_template_fallback(str(exc))

    return _render_blank_report(
        output_path=output_path,
        template_instructions=template_instructions,
        metrics=metrics,
        payment_profile=payment_profile,
        duration_profile=duration_profile,
        report_plan=report_plan,
        charts=charts,
    )


def _render_blank_report(
    *,
    output_path: Path,
    template_instructions: TemplateInstructions,
    metrics: ParkingMetrics,
    payment_profile: PaymentProfile,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
) -> Path:
    """当模板无法识别时，从空白文档生成标准报告作为兜底。"""
    document = Document()
    _configure_styles(document)

    title = template_instructions.title or "停车明细分析报告"
    _add_title(document, title)
    _add_report_meta(document, report_plan)
    _add_metrics_table(document, metrics)
    _add_payment_section(document, payment_profile, report_plan)
    _add_duration_section(document, duration_profile, report_plan)
    _add_charts_section(document, report_plan, charts)
    _add_observations_section(document, report_plan)
    _add_recommendations_section(document, report_plan)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _render_template_aware_report(
    *,
    output_path: Path,
    template_path: Path,
    metrics: ParkingMetrics,
    payment_profile: PaymentProfile,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
    data_period: str | None,
    generated_at: datetime | None,
) -> Path:
    """打开原 DOCX 模板，在原位置替换占位并尽量保留模板样式。"""
    document = Document(str(template_path))

    _fill_report_info(document, data_period, generated_at)
    _fill_metrics_table(document, metrics)
    _replace_template_paragraphs(document, payment_profile, duration_profile, report_plan, charts)
    _remove_template_instruction_banner(document)
    _format_document_runs(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def _configure_styles(document: Document) -> None:
    """设置报告的基础字体和段落样式。"""
    normal = document.styles["Normal"]
    _set_style_font(normal, BODY_FONT)
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        _set_style_font(style, BODY_FONT)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        _set_style_font(style, HEADING_FONT)
        style.font.bold = True


def _add_title(document: Document, title: str) -> None:
    """写入报告标题。"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    _set_run_font(run, HEADING_FONT)
    document.add_paragraph()


def _add_report_meta(document: Document, report_plan: ReportPlan) -> None:
    """写入报告生成模式等元信息，便于面试官验证 agent 行为。"""
    paragraph = document.add_paragraph()
    label = paragraph.add_run("生成模式：")
    label.bold = True
    _set_run_font(label, BODY_FONT)
    value = paragraph.add_run(f"{report_plan.planner_mode}；fallback={report_plan.fallback_used}")
    _set_run_font(value, BODY_FONT)


def _add_metrics_table(document: Document, metrics: ParkingMetrics) -> None:
    """写入六个硬指标，确保 Agent 不参与计算。"""
    document.add_heading("一、关键指标", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "数值"
    top_payment = (
        f"{metrics.top_payment_method.method}（{metrics.top_payment_method.count} 笔）"
        if metrics.top_payment_method
        else "N/A"
    )
    rows = [
        ("交易总笔数", str(metrics.total_transactions)),
        ("应收总金额", _format_decimal(metrics.total_receivable)),
        ("实收总金额", _format_decimal(metrics.total_collected)),
        ("实际抵扣总额", _format_decimal(metrics.total_actual_deductions)),
        ("实收率", f"{metrics.collection_rate_pct}%"),
        ("最高频支付方式", top_payment),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    _format_table_font(table)


def _add_payment_section(
    document: Document,
    payment_profile: PaymentProfile,
    report_plan: ReportPlan,
) -> None:
    """写入支付方式与支付渠道摘要。"""
    document.add_heading("二、支付方式与渠道", level=1)
    document.add_paragraph(report_plan.payment_section_summary)
    method = payment_profile.dominant_payment_method
    channel = payment_profile.dominant_payment_channel
    if method:
        document.add_paragraph(
            f"最高频支付方式为 {method.name}，共 {method.count} 笔，占比 {method.count_pct}%。"
        )
    if channel:
        document.add_paragraph(
            f"最高频支付渠道为 {channel.name}，共 {channel.count} 笔，占比 {channel.count_pct}%。"
        )


def _add_duration_section(
    document: Document,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
) -> None:
    """写入停车时长分析摘要。"""
    document.add_heading("三、停车时长分析", level=1)
    document.add_paragraph(report_plan.duration_section_summary)
    document.add_paragraph(
        "平均停车时长 "
        f"{duration_profile.average_hours} 小时，中位数 {duration_profile.median_hours} 小时，"
        f"最长 {duration_profile.max_hours} 小时，超过 12 小时记录 {duration_profile.over_12h_count} 笔。"
    )


def _add_charts_section(
    document: Document,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
) -> None:
    """插入 Agent 选择的真实图表；若选择缺失则回退到第一张可用图表。"""
    document.add_heading("四、图表分析", level=1)
    selected_ids = [chart.chart_id for chart in report_plan.selected_charts if chart.chart_id in charts]
    if not selected_ids and charts:
        selected_ids = [next(iter(charts))]

    for chart_id in selected_ids:
        chart = charts[chart_id]
        document.add_paragraph(chart.title)
        document.add_picture(str(chart.path), width=Inches(5.8))


def _add_observations_section(document: Document, report_plan: ReportPlan) -> None:
    """写入 Agent 选择的补充观察。"""
    document.add_heading("五、补充观察", level=1)
    for observation in report_plan.observations:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(observation.title + "：").bold = True
        paragraph.add_run(observation.text)
        document.add_paragraph(f"经营含义：{observation.business_implication}")


def _add_recommendations_section(document: Document, report_plan: ReportPlan) -> None:
    """写入结论与建议。"""
    document.add_heading("六、结论与建议", level=1)
    for recommendation in report_plan.recommendations:
        document.add_paragraph(recommendation.text, style="List Number")


def _fill_report_info(document: Document, data_period: str | None, generated_at: datetime | None) -> None:
    """替换模板报告信息中的数据周期和生成时间占位。"""
    generated_text = (generated_at or datetime.now()).strftime("%Y-%m-%d %H:%M:%S")
    period_text = data_period or "详见上传数据"

    for paragraph in document.paragraphs:
        text = paragraph.text
        if "数据周期：" in text:
            _replace_paragraph_text(paragraph, f"数据周期：{period_text}")
        elif "生成时间：" in text:
            _replace_paragraph_text(paragraph, f"生成时间：{generated_text}")


def _fill_metrics_table(document: Document, metrics: ParkingMetrics) -> None:
    """填写关键指标，并删除仅供模板说明使用的示例预期值列。"""
    table = _find_metrics_table(document)
    if table is None:
        raise ValueError("Template metrics table not found")

    values = _metrics_values(metrics)
    if len(table.rows) < len(values) + 1 or len(table.columns) < 3:
        raise ValueError("Template metrics table shape is invalid")

    for row_index, value in enumerate(values, start=1):
        _set_cell_text(table.rows[row_index].cells[1], value)

    _remove_table_column(table, 2)


def _remove_table_column(table, column_index: int) -> None:
    """从 Word 表格行和网格定义中完整删除指定列。"""
    for row in table.rows:
        cells = row._tr.tc_lst
        if column_index >= len(cells):
            raise ValueError("Template metrics table rows have inconsistent columns")
        row._tr.remove(cells[column_index])

    grid_columns = table._tbl.tblGrid.gridCol_lst
    if column_index >= len(grid_columns):
        raise ValueError("Template metrics table grid is inconsistent")
    table._tbl.tblGrid.remove(grid_columns[column_index])


def _find_metrics_table(document: Document):
    """按表头识别关键指标表格，避免依赖固定表格序号。"""
    for table in document.tables:
        if len(table.rows) == 0 or len(table.columns) < 3:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells[:3]]
        if headers == ["指标", "填写值", "示例预期值"]:
            return table
    return None


def _metrics_values(metrics: ParkingMetrics) -> list[str]:
    """生成关键指标表格第二列的确定性填写值。"""
    top_payment = (
        f"{metrics.top_payment_method.method} ×{metrics.top_payment_method.count}"
        if metrics.top_payment_method
        else "N/A"
    )
    return [
        str(metrics.total_transactions),
        f"¥{_format_decimal(metrics.total_receivable)}",
        f"¥{_format_decimal(metrics.total_collected)}",
        f"¥{_format_decimal(metrics.total_actual_deductions)}",
        f"{metrics.collection_rate_pct}%",
        top_payment,
    ]


def _replace_template_paragraphs(
    document: Document,
    payment_profile: PaymentProfile,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
) -> None:
    """按模板占位文本替换章节内容，并在图表占位附近插入真实图表。"""
    paragraphs = list(document.paragraphs)
    for paragraph in paragraphs:
        text = paragraph.text
        if "【叙述：各支付方式与支付渠道" in text:
            _replace_paragraph_text(paragraph, report_plan.payment_section_summary)
            _append_payment_detail_after(paragraph, payment_profile)
        elif "【占位图" in text:
            _replace_chart_placeholder(paragraph, report_plan, charts)
        elif text.startswith("【根据“收费时间"):
            _replace_duration_placeholder(paragraph, duration_profile, report_plan)
        elif (
            "停车时长分布柱形图" in text
            or "每日入场时间分布图" in text
            or "每日收费时间分布柱形图" in text
            or text.strip() == "】"
        ):
            _remove_paragraph(paragraph)
        elif "【智能体生成的观察内容置于此处" in text:
            _replace_observations_placeholder(paragraph, report_plan)
        elif "读者对象：" in text:
            _remove_paragraph(paragraph)
        elif "【依据上述分析给出" in text:
            _replace_recommendations_placeholder(paragraph, report_plan)


def _append_payment_detail_after(paragraph: Paragraph, payment_profile: PaymentProfile) -> None:
    """在支付摘要后补充支付方式和渠道的确定性明细。"""
    current = paragraph
    method = payment_profile.dominant_payment_method
    channel = payment_profile.dominant_payment_channel
    if method:
        current = _insert_paragraph_after(
            current,
            f"最高频支付方式为 {method.name}，共 {method.count} 笔，占比 {method.count_pct}%。",
        )
    if channel:
        _insert_paragraph_after(
            current,
            f"最高频支付渠道为 {channel.name}，共 {channel.count} 笔，占比 {channel.count_pct}%。",
        )


def _replace_chart_placeholder(
    paragraph: Paragraph,
    report_plan: ReportPlan,
    charts: dict[str, RenderedChart],
) -> None:
    """把模板图表占位替换为 Agent 选择的真实图表。"""
    _remove_adjacent_placeholder_images(paragraph)
    selected_ids = [chart.chart_id for chart in report_plan.selected_charts if chart.chart_id in charts]
    if not selected_ids and charts:
        selected_ids = [next(iter(charts))]
    if not selected_ids:
        _replace_paragraph_text(paragraph, "无可用图表。")
        return

    first_chart = charts[selected_ids[0]]
    _replace_paragraph_text(paragraph, first_chart.title)
    current = paragraph
    for chart_id in selected_ids:
        chart = charts[chart_id]
        picture_paragraph = _insert_paragraph_after(current, "")
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = picture_paragraph.add_run()
        run.add_picture(str(chart.path), width=Inches(5.8))
        current = picture_paragraph


def _replace_duration_placeholder(
    paragraph: Paragraph,
    duration_profile: DurationProfile,
    report_plan: ReportPlan,
) -> None:
    """替换停车时长分析占位，并移除模板中紧随的图表说明占位行。"""
    _replace_paragraph_text(paragraph, report_plan.duration_section_summary)
    _insert_paragraph_after(
        paragraph,
        "平均停车时长 "
        f"{duration_profile.average_hours} 小时，中位数 {duration_profile.median_hours} 小时，"
        f"最长 {duration_profile.max_hours} 小时，超过 12 小时记录 {duration_profile.over_12h_count} 笔。",
    )


def _replace_observations_placeholder(paragraph: Paragraph, report_plan: ReportPlan) -> None:
    """将补充观察占位替换为 Agent 输出的 2-3 条观察。"""
    lines = [
        f"{index}. {observation.title}：{observation.text} 经营含义：{observation.business_implication}"
        for index, observation in enumerate(report_plan.observations, start=1)
    ]
    _replace_with_lines(paragraph, lines)


def _replace_recommendations_placeholder(paragraph: Paragraph, report_plan: ReportPlan) -> None:
    """将结论与建议占位替换为 Agent 输出的 2-4 条建议。"""
    lines = [
        f"{index}. {recommendation.text}"
        for index, recommendation in enumerate(report_plan.recommendations, start=1)
    ]
    _replace_with_lines(paragraph, lines)


def _remove_template_instruction_banner(document: Document) -> None:
    """移除模板顶部说明，避免最终报告保留“红色斜体占位”提示。"""
    for paragraph in list(document.paragraphs):
        if "报告模板" in paragraph.text and "占位内容" in paragraph.text:
            _remove_paragraph(paragraph)


def _remove_adjacent_placeholder_images(paragraph: Paragraph) -> None:
    """删除图表占位说明附近的模板示例图片，避免最终报告图表重复。"""
    for candidate in [paragraph._p.getprevious(), paragraph._p.getnext()]:
        if candidate is not None and _xml_contains_drawing(candidate):
            parent = candidate.getparent()
            if parent is not None:
                parent.remove(candidate)


def _xml_contains_drawing(element) -> bool:
    """判断 OOXML 段落是否包含图片或绘图对象。"""
    return any(child.tag.endswith("}drawing") or child.tag.endswith("}pict") for child in element.iter())


def _format_decimal(value: Decimal) -> str:
    """统一格式化金额类 Decimal，避免科学计数法或多余精度。"""
    return f"{value:.2f}"


def _set_style_font(style, font_name: str) -> None:
    """同时设置西文字体和中文 East Asia 字体，避免 Word 自动 fallback。"""
    style.font.name = font_name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_run_font(run, font_name: str) -> None:
    """设置 run 的中英文字体；用于标题、加粗片段和表格内容。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _format_table_font(table) -> None:
    """统一表格内所有单元格文字字体，避免表格内容与正文风格不一致。"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, BODY_FONT)


def _format_document_runs(document: Document) -> None:
    """补充设置中文 East Asia 字体，同时尽量不改动模板原有段落样式。"""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, BODY_FONT)
    for table in document.tables:
        _format_table_font(table)


def _set_cell_text(cell, text: str) -> None:
    """替换单元格文本并保留单元格和表格样式。"""
    paragraph = cell.paragraphs[0]
    _replace_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        _remove_paragraph(extra)


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """替换段落文本，保留段落级样式并统一中文字体。"""
    paragraph.clear()
    run = paragraph.add_run(text)
    _set_run_font(run, BODY_FONT)


def _replace_with_lines(paragraph: Paragraph, lines: list[str]) -> None:
    """用多行内容替换一个占位段落，后续行插入在原段落之后。"""
    if not lines:
        _replace_paragraph_text(paragraph, "")
        return

    _replace_paragraph_text(paragraph, lines[0])
    current = paragraph
    for line in lines[1:]:
        current = _insert_paragraph_after(current, line)


def _insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    """在指定段落后插入新段落，并沿用原段落样式。"""
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if paragraph.style is not None:
        new_paragraph.style = paragraph.style
    if text:
        run = new_paragraph.add_run(text)
        _set_run_font(run, BODY_FONT)
    return new_paragraph


def _remove_paragraph(paragraph: Paragraph) -> None:
    """从文档中移除占位段落。"""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
